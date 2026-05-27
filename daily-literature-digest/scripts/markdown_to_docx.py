#!/usr/bin/env python3
"""Convert a daily literature digest Markdown file to a polished DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = "1F4E5F"
ACCENT_DARK = "173E4F"
ACCENT_LIGHT = "EAF4F5"
BORDER = "B8CDD2"
TEXT = RGBColor(35, 40, 45)
MUTED = RGBColor(95, 105, 115)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size: float = 10.5, color: RGBColor = TEXT) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin_twips: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, header: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text.strip())
    set_run_font(run, size=9.5 if not header else 9.3, bold=header, color=RGBColor(255, 255, 255) if header else TEXT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_border(cell)
    if header:
        shade_cell(cell, ACCENT)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows = []
    for line in lines:
        parts = [part.strip() for part in line.strip().strip("|").split("|")]
        rows.append(parts)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", part.replace(" ", "")) for part in rows[1]):
        rows.pop(1)
    header = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    return header, body


def add_table(document: Document, lines: list[str]) -> None:
    header, body = parse_table(lines)
    if not header:
        return
    table = document.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, heading in enumerate(header):
        set_cell_text(table.rows[0].cells[index], heading, header=True, center=index < 3)
    for row in body:
        cells = table.add_row().cells
        for index, value in enumerate(row[: len(header)]):
            set_cell_text(cells[index], value, center=index < 3)
        for index in range(len(row), len(header)):
            set_cell_text(cells[index], "", center=index < 3)
    document.add_paragraph()


def add_paragraph_with_inline_bold(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.16
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=10.5, bold=True, color=TEXT)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=10.5, color=TEXT)
    return paragraph


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color in [
        ("Title", 20, ACCENT_DARK),
        ("Heading 1", 16, ACCENT_DARK),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11.5, ACCENT_DARK),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name == "Title" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Daily Literature Digest")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated for Hanyu Zheng")
    set_run_font(run, size=8.5, color=MUTED)


def add_note_box(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    table = document.add_table(rows=len(lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, line in zip(table.rows, lines):
        cell = row.cells[0]
        set_cell_text(cell, line.lstrip("- ").strip())
        shade_cell(cell, ACCENT_LIGHT)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9.8, color=TEXT)
    document.add_paragraph()


def convert_markdown_to_docx(markdown_path: Path, docx_path: Path) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    document = Document()
    style_document(document)
    add_header_footer(document)

    pending_meta: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            if pending_meta:
                add_note_box(document, pending_meta)
                pending_meta = []
            add_table(document, table_lines)
            continue

        if stripped.startswith("# "):
            if pending_meta:
                add_note_box(document, pending_meta)
                pending_meta = []
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(stripped[2:].strip())
            set_run_font(run, size=20, bold=True, color=RGBColor.from_string(ACCENT_DARK))
        elif stripped.startswith("## "):
            if pending_meta:
                add_note_box(document, pending_meta)
                pending_meta = []
            paragraph = document.add_paragraph(stripped[3:].strip(), style="Heading 1")
            set_paragraph_font(paragraph, size=16, color=RGBColor.from_string(ACCENT_DARK))
        elif stripped.startswith("### "):
            if pending_meta:
                add_note_box(document, pending_meta)
                pending_meta = []
            paragraph = document.add_paragraph(stripped[4:].strip(), style="Heading 2")
            set_paragraph_font(paragraph, size=13, color=RGBColor.from_string(ACCENT))
        elif stripped.startswith("- "):
            if stripped.startswith(("- 收件人", "- 检索窗口", "- 数据源", "- 本次", "- 说明")):
                pending_meta.append(stripped)
            else:
                if pending_meta:
                    add_note_box(document, pending_meta)
                    pending_meta = []
                paragraph = add_paragraph_with_inline_bold(document, stripped[2:].strip(), style="List Bullet")
                set_paragraph_font(paragraph, size=10.3)
        else:
            if pending_meta:
                add_note_box(document, pending_meta)
                pending_meta = []
            add_paragraph_with_inline_bold(document, stripped)
        index += 1

    if pending_meta:
        add_note_box(document, pending_meta)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = lines[0].lstrip("# ").strip() if lines else markdown_path.stem
    document.core_properties.author = "Codex"
    document.save(docx_path)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert a daily literature digest Markdown file to DOCX.")
    parser.add_argument("markdown_path")
    parser.add_argument("docx_path")
    return parser


def main() -> int:
    args = build_parser().parse_args()
    convert_markdown_to_docx(Path(args.markdown_path), Path(args.docx_path))
    print(str(Path(args.docx_path).resolve()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
